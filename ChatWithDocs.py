import streamlit as st
import os
import io
import docx
from dotenv import load_dotenv
from langchain_core.messages import AIMessage, HumanMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_cohere import CohereEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from heading import get_heading
from PyPDF2 import PdfReader
from langchain.docstore.document import Document

get_heading("Chat With Documents")


load_dotenv()
google_api_key = st.session_state.google_api_key




google_models = [
    "gemini-1.5-flash",
    "gemini-1.5-pro",
]

def reset_conversation():
    keys_to_reset = [
        "chat_history_docs",
        "uploaded_files_content",
        "files_processed",
        "vectorstore"
    ]
    for key in keys_to_reset:
        if key in st.session_state:
            del st.session_state[key]
    st.session_state.chat_history_docs = [
        AIMessage(content="Hi there! How can I help you today?"),
    ]


def get_pdf_text(pdf_file):
    text = ""
    # for pdf in pdf_file:
    pdf_reader = PdfReader(pdf_file)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def get_docx_text(docx_file):
    text = ""
    doc = docx.Document(docx_file)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)
    text += " ".join(full_text)

    return text


def get_documents_text(doc):

    split_tup = os.path.splitext(doc.name)
    file_extension = split_tup[1]

    if file_extension == ".docx":
        text = get_docx_text(doc)

    elif file_extension == ".pdf":
        text = get_pdf_text(doc)
    

    return text


def get_text_chunks(text, filename):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size= 800,
        chunk_overlap= 150,
        length_function=len
    )
    chunks = text_splitter.split_text(text)

    doc_list = []
    for chunk in chunks:
        metadata = {"source": filename}
        doc_string = Document(page_content=chunk, metadata=metadata)
        doc_list.append(doc_string)
    return doc_list


def get_vectorstore(text_chunks_list):
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    vectorstore = FAISS.from_documents(documents=text_chunks_list, embedding=embeddings)
    return vectorstore


def get_context_retriever_chain(vectorstore, llm):
    
    retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

    contextualize_q_system_prompt = """Given a chat history and the latest user question 
    which might reference context in the chat history, formulate a standalone question 
    which can be understood without the chat history. Do NOT answer the question, 
    just reformulate it if needed and otherwise return it as is. Do NOT give the source."""

    contextualize_q_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", contextualize_q_system_prompt),
            MessagesPlaceholder("chat_history_docs"),
            ("human", "{input}"),
        ]
    )
    history_aware_retriever = create_history_aware_retriever(
        llm, retriever, contextualize_q_prompt
    )
    return history_aware_retriever


def get_conversational_rag_chain(history_aware_retriever, llm):

    qa_system_prompt = """You are a helpful assistant for question-answering tasks. 
    Use the following pieces of retrieved context to answer the question. Do not mention
    any source for the documents.
    If you don't know the answer, just say that you don't know. 
    Keep the answer concise.\n\n

    {context}"""

    qa_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", qa_system_prompt),
            MessagesPlaceholder("chat_history_docs"),
            ("human", "{input}"),
        ]
    )
    question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)

    rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)

    return rag_chain


def get_response(user_question, llm):
    retriever_chain = get_context_retriever_chain(st.session_state.vectorstore, llm)
    conversational_rag_chain = get_conversational_rag_chain(retriever_chain, llm)

    response = conversational_rag_chain.invoke({
        "chat_history_docs": st.session_state.chat_history_docs,
        "input": user_question
    })
    
    source = response['context'][0].metadata['source']
    # Create the final response
    final_response = f"{response['answer']} \n\n **Source Document:** `{source}`"

    return final_response


def get_llm_model(model_name, temperature, google_api_key, max_output_tokens):

    llm = ChatGoogleGenerativeAI(model=model_name,
                                 google_api_key=google_api_key,
                                 temperature=temperature,
                                 max_output_tokens=max_output_tokens)
    return llm


if "chat_history_docs" not in st.session_state:
        st.session_state.chat_history_docs = [
            AIMessage(content="Hi there! How can I help you today?"),
        ]

if "uploaded_files_content" not in st.session_state:
    st.session_state.uploaded_files_content = []

if "files_processed" not in st.session_state:
    st.session_state.files_processed = False

if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None



with st.sidebar:
    with st.popover("Upload your documents 🔗", use_container_width=True):
        uploaded_files = st.file_uploader("You can upload multiple files", type=["docx", "pdf"], accept_multiple_files=True)
    
    if uploaded_files:
        # Store the file contents
        new_files = [file for file in uploaded_files if file not in st.session_state.uploaded_files_content]
        for file in new_files:
            file_content = file.read()
            st.session_state.uploaded_files_content.append({"name": file.name, "content": file_content, "type": file.type})
        
        if st.button("Process", use_container_width=True, type="primary"):
            text_chunks_list = []
            with st.spinner("Processing..."):
                try:
                    for file_data in st.session_state.uploaded_files_content:
                        filename = file_data["name"]
                        file_content = io.BytesIO(file_data["content"])
                        file_content.name = filename  # Set the name attribute for compatibility
                        text = get_documents_text(file_content)
                        text_chunks = get_text_chunks(text, filename)
                        text_chunks_list.extend(text_chunks)
                    st.session_state.vectorstore = get_vectorstore(text_chunks_list)
                    st.session_state.files_processed = True
                except Exception as e:
                    st.write(f"An Error has occurred \n\n{e}")



    if st.session_state.files_processed:
        model_name = st.selectbox("Select a Model", options=google_models)
        
        col1, col2 = st.columns(2)
        with col1:
            with st.popover("⚙️Parameters"):
                temperature = st.slider("Temperature", min_value=0.0, max_value=1.0, value=0.3, step=0.1)
                max_output_tokens = st.slider("Max Tokens", min_value=200, max_value=2000, value=1000, step=200)
        
        with col2:
            st.button(
                "🗑️ Reset", 
                on_click=reset_conversation,
                use_container_width=True
                )



# Display chat history
for message in st.session_state.chat_history_docs:
    if isinstance(message, AIMessage):
        with st.chat_message("AI", avatar="bot.png"):
            st.write(message.content)
    elif isinstance(message, HumanMessage):
        with st.chat_message("Human", avatar="human.png"):
            st.write(message.content)


if st.session_state.files_processed:
    llm = get_llm_model(model_name, temperature, google_api_key, max_output_tokens)

    if user_query := st.chat_input("Type your question..."):
        st.session_state.chat_history_docs.append(HumanMessage(content=user_query))

        with st.chat_message("Human", avatar="human.png"):
            st.write(user_query)

        with st.spinner("Thinking..."):
            try:
                response = get_response(user_query, llm)
            
                with st.chat_message("AI", avatar="bot.png"):
                    st.write(response)
            
            except Exception as e:
                st.write(f"An Error has occurred \n\n{e}")
                
        st.session_state.chat_history_docs.append(AIMessage(content=response))
else:
    st.info("Please upload and process the files to start chatting.")







